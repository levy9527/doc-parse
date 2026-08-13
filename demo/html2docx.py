from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from bs4 import BeautifulSoup

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_paragraph_border(paragraph, color="3157d5", size=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="8" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def html_to_docx(html_path, docx_path):
    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Cm(0.64)
    r = style.element.rPr
    if r is None:
        r = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(r)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>'
    )
    r.append(rFonts)

    for h_level in range(1, 5):
        heading_style = doc.styles[f"Heading {h_level}"]
        heading_style.font.name = "微软雅黑"
        heading_style.font.color.rgb = RGBColor(0x22, 0x3A, 0x75)
        hr = heading_style.element.rPr
        if hr is None:
            hr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            heading_style.element.append(hr)
        hFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>'
        )
        hr.append(hFonts)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
        heading_style.paragraph_format.first_line_indent = Cm(0)
        if h_level == 1:
            heading_style.font.size = Pt(24)
        elif h_level == 2:
            heading_style.font.size = Pt(16)
        elif h_level == 3:
            heading_style.font.size = Pt(14)

    article = soup.find("article")
    if not article:
        article = soup

    for el in article.children:
        if isinstance(el, str):
            continue

        tag = el.name
        if tag is None:
            continue

        if tag == "h1":
            p = doc.add_heading(el.get_text(strip=True), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif tag == "p":
            cls = el.get("class", [])
            if "subtitle" in cls:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(el.get_text(strip=True))
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x66, 0x71, 0x8A)
                run.font.name = "微软雅黑"
            else:
                doc.add_paragraph(el.get_text(strip=True))

        elif tag == "blockquote":
            text = el.get_text(strip=True)
            p = doc.add_paragraph(text)
            add_paragraph_border(p)
            p.paragraph_format.left_indent = Cm(1)

        elif tag in ("h2", "h3", "h4"):
            level = int(tag[1])
            p = doc.add_heading(el.get_text(strip=True), level=level)

        elif tag == "ul":
            for li in el.find_all("li", recursive=False):
                text = li.get_text(strip=True)
                p = doc.add_paragraph(text, style="List Bullet")
                p.paragraph_format.first_line_indent = Cm(0)

        elif tag == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            col_count = max(len(r.find_all(["th", "td"])) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for j, cell in enumerate(cells):
                    cell_text = cell.get_text(strip=True)
                    tbl_cell = table.cell(i, j)
                    tbl_cell.text = cell_text
                    for par in tbl_cell.paragraphs:
                        par.paragraph_format.first_line_indent = Cm(0)
                        for run in par.runs:
                            run.font.size = Pt(10)
                            run.font.name = "微软雅黑"
                            if cell.name == "th":
                                run.font.bold = True
                    if cell.name == "th":
                        set_cell_shading(tbl_cell, "EEF2FF")

        elif tag == "section":
            cls = el.get("class", [])
            if "sources" in cls:
                for child in el.children:
                    if isinstance(child, str):
                        continue
                    if child.name == "h2":
                        doc.add_heading(child.get_text(strip=True), level=2)
                    elif child.name == "ul":
                        for li in child.find_all("li", recursive=False):
                            p = doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
                            p.paragraph_format.first_line_indent = Cm(0)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == "__main__":
    html_path = "/Users/levy/Downloads/2026年3月金融行业网络安全监测与态势研判报告.html"
    docx_path = "/Users/levy/PycharmProjects/doc-parse/2026年3月金融行业网络安全监测与态势研判报告.docx"
    html_to_docx(html_path, docx_path)
